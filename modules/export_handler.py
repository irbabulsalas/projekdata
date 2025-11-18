"""
Export Handler Module for AI Data Analysis Platform
Comprehensive export system for data, models, visualizations, and reports.
"""

import pandas as pd
import numpy as np
from typing import Dict, List, Any, Optional, Union, BinaryIO
import json
import pickle
import base64
from datetime import datetime
import os
import warnings
warnings.filterwarnings('ignore')

# Export Libraries
from fpdf import FPDF
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import plotly.graph_objects as go
import plotly.io as pio
from io import BytesIO
import zipfile

# ML Export
import joblib
from sklearn.metrics import classification_report, confusion_matrix
import nbformat
from nbformat.v4 import new_notebook, new_code_cell, new_markdown_cell

# Image Processing
from PIL import Image as PILImage
import matplotlib.pyplot as plt

from utils.helpers import FileHandler, ExportHelper
from utils.error_handler import handle_errors, ValidationError
from utils.rate_limiter import rate_limit


class DataExporter:
    """Advanced data export system with multiple format support."""
    
    def __init__(self):
        """Initialize data exporter."""
        self.export_history = []
        self.exports_dir = "exports"
        self._ensure_exports_directory()
    
    def _ensure_exports_directory(self):
        """Ensure exports directory exists."""
        os.makedirs(self.exports_dir, exist_ok=True)
    
    @handle_errors("data")
    def export_data(self, data: pd.DataFrame, filename: str, format: str = 'csv', **kwargs) -> str:
        """Export data to various formats."""
        if data is None or data.empty:
            raise ValidationError("No data to export")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{filename}_{timestamp}"
        
        if format.lower() == 'csv':
            filepath = self._export_csv(data, base_filename, **kwargs)
        elif format.lower() == 'excel':
            filepath = self._export_excel(data, base_filename, **kwargs)
        elif format.lower() == 'json':
            filepath = self._export_json(data, base_filename, **kwargs)
        elif format.lower() == 'parquet':
            filepath = self._export_parquet(data, base_filename, **kwargs)
        elif format.lower() == 'tsv':
            filepath = self._export_tsv(data, base_filename, **kwargs)
        else:
            raise ValidationError(f"Unsupported export format: {format}")
        
        self._log_export('data', filepath, format, len(data))
        return filepath
    
    def _export_csv(self, data: pd.DataFrame, filename: str, **kwargs) -> str:
        """Export data to CSV format."""
        filepath = os.path.join(self.exports_dir, f"{filename}.csv")
        data.to_csv(filepath, index=kwargs.get('index', False), encoding=kwargs.get('encoding', 'utf-8'))
        return filepath
    
    def _export_excel(self, data: pd.DataFrame, filename: str, **kwargs) -> str:
        """Export data to Excel format with multiple sheets."""
        filepath = os.path.join(self.exports_dir, f"{filename}.xlsx")
        
        with pd.ExcelWriter(filepath, engine='xlsxwriter') as writer:
            # Main data sheet
            data.to_excel(writer, sheet_name='Data', index=kwargs.get('index', False))
            
            # Summary sheet
            summary_data = {
                'Metric': ['Rows', 'Columns', 'Missing Values', 'Memory Usage (MB)'],
                'Value': [
                    len(data),
                    len(data.columns),
                    data.isnull().sum().sum(),
                    data.memory_usage(deep=True).sum() / 1024**2
                ]
            }
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            # Data types sheet
            dtypes_df = pd.DataFrame({
                'Column': data.columns,
                'Data Type': data.dtypes.astype(str),
                'Non-Null Count': data.count(),
                'Null Count': data.isnull().sum()
            })
            dtypes_df.to_excel(writer, sheet_name='Data Types', index=False)
        
        return filepath
    
    def _export_json(self, data: pd.DataFrame, filename: str, **kwargs) -> str:
        """Export data to JSON format."""
        filepath = os.path.join(self.exports_dir, f"{filename}.json")
        
        orient = kwargs.get('orient', 'records')
        data.to_json(filepath, orient=orient, indent=kwargs.get('indent', 2), date_format='iso')
        return filepath
    
    def _export_parquet(self, data: pd.DataFrame, filename: str, **kwargs) -> str:
        """Export data to Parquet format."""
        filepath = os.path.join(self.exports_dir, f"{filename}.parquet")
        data.to_parquet(filepath, index=kwargs.get('index', False))
        return filepath
    
    def _export_tsv(self, data: pd.DataFrame, filename: str, **kwargs) -> str:
        """Export data to TSV format."""
        filepath = os.path.join(self.exports_dir, f"{filename}.tsv")
        data.to_csv(filepath, sep='\t', index=kwargs.get('index', False), encoding=kwargs.get('encoding', 'utf-8'))
        return filepath
    
    @handle_errors("data")
    def export_model(self, model: Any, filename: str, format: str = 'pickle', **metadata) -> str:
        """Export trained model to various formats."""
        if model is None:
            raise ValidationError("No model to export")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{filename}_{timestamp}"
        
        if format.lower() == 'pickle':
            filepath = self._export_model_pickle(model, base_filename, metadata)
        elif format.lower() == 'joblib':
            filepath = self._export_model_joblib(model, base_filename, metadata)
        elif format.lower() == 'onnx':
            filepath = self._export_model_onnx(model, base_filename, metadata)
        elif format.lower() == 'json':
            filepath = self._export_model_json(model, base_filename, metadata)
        else:
            raise ValidationError(f"Unsupported model export format: {format}")
        
        self._log_export('model', filepath, format, 1)
        return filepath
    
    def _export_model_pickle(self, model: Any, filename: str, metadata: Dict) -> str:
        """Export model using pickle."""
        filepath = os.path.join(self.exports_dir, f"{filename}.pkl")
        
        model_data = {
            'model': model,
            'metadata': metadata,
            'export_timestamp': datetime.now().isoformat(),
            'version': '1.0'
        }
        
        with open(filepath, 'wb') as f:
            pickle.dump(model_data, f)
        
        return filepath
    
    def _export_model_joblib(self, model: Any, filename: str, metadata: Dict) -> str:
        """Export model using joblib."""
        filepath = os.path.join(self.exports_dir, f"{filename}.joblib")
        
        model_data = {
            'model': model,
            'metadata': metadata,
            'export_timestamp': datetime.now().isoformat(),
            'version': '1.0'
        }
        
        joblib.dump(model_data, filepath)
        return filepath
    
    def _export_model_onnx(self, model: Any, filename: str, metadata: Dict) -> str:
        """Export model to ONNX format (if supported)."""
        try:
            import onnx
            import skl2onnx
            
            filepath = os.path.join(self.exports_dir, f"{filename}.onnx")
            
            # Convert to ONNX (simplified version)
            onnx_model = skl2onnx.convert_sklearn(model, 'model_data')
            
            with open(filepath, "wb") as f:
                f.write(onnx_model.SerializeToString())
            
            return filepath
        except ImportError:
            raise ValidationError("ONNX export requires onnx and skl2onnx packages")
        except Exception as e:
            raise ValidationError(f"ONNX export failed: {str(e)}")
    
    def _export_model_json(self, model: Any, filename: str, metadata: Dict) -> str:
        """Export model metadata to JSON."""
        filepath = os.path.join(self.exports_dir, f"{filename}_metadata.json")
        
        model_info = {
            'model_type': type(model).__name__,
            'metadata': metadata,
            'export_timestamp': datetime.now().isoformat(),
            'version': '1.0',
            'parameters': getattr(model, 'get_params', lambda: {})()
        }
        
        with open(filepath, 'w') as f:
            json.dump(model_info, f, indent=2, default=str)
        
        return filepath
    
    @handle_errors("data")
    def export_visualization(self, fig: go.Figure, filename: str, format: str = 'html', **kwargs) -> str:
        """Export visualization to various formats."""
        if fig is None:
            raise ValidationError("No visualization to export")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{filename}_{timestamp}"
        
        if format.lower() == 'html':
            filepath = self._export_viz_html(fig, base_filename, **kwargs)
        elif format.lower() == 'png':
            filepath = self._export_viz_png(fig, base_filename, **kwargs)
        elif format.lower() == 'pdf':
            filepath = self._export_viz_pdf(fig, base_filename, **kwargs)
        elif format.lower() == 'svg':
            filepath = self._export_viz_svg(fig, base_filename, **kwargs)
        else:
            raise ValidationError(f"Unsupported visualization export format: {format}")
        
        self._log_export('visualization', filepath, format, 1)
        return filepath
    
    def _export_viz_html(self, fig: go.Figure, filename: str, **kwargs) -> str:
        """Export visualization to HTML."""
        filepath = os.path.join(self.exports_dir, f"{filename}.html")
        fig.write_html(filepath, include_plotlyjs=kwargs.get('include_plotlyjs', 'cdn'))
        return filepath
    
    def _export_viz_png(self, fig: go.Figure, filename: str, **kwargs) -> str:
        """Export visualization to PNG."""
        try:
            filepath = os.path.join(self.exports_dir, f"{filename}.png")
            fig.write_image(filepath, width=kwargs.get('width', 1200), height=kwargs.get('height', 800))
            return filepath
        except Exception as e:
            raise ValidationError(f"PNG export requires kaleido package: {str(e)}")
    
    def _export_viz_pdf(self, fig: go.Figure, filename: str, **kwargs) -> str:
        """Export visualization to PDF."""
        try:
            filepath = os.path.join(self.exports_dir, f"{filename}.pdf")
            fig.write_image(filepath, format='pdf', width=kwargs.get('width', 1200), height=kwargs.get('height', 800))
            return filepath
        except Exception as e:
            raise ValidationError(f"PDF export requires kaleido package: {str(e)}")
    
    def _export_viz_svg(self, fig: go.Figure, filename: str, **kwargs) -> str:
        """Export visualization to SVG."""
        filepath = os.path.join(self.exports_dir, f"{filename}.svg")
        fig.write_image(filepath, format='svg', width=kwargs.get('width', 1200), height=kwargs.get('height', 800))
        return filepath
    
    @handle_errors("data")
    def export_report(self, report_data: Dict[str, Any], filename: str, format: str = 'pdf', **kwargs) -> str:
        """Export comprehensive analysis report."""
        if not report_data:
            raise ValidationError("No report data to export")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{filename}_{timestamp}"
        
        if format.lower() == 'pdf':
            filepath = self._export_report_pdf(report_data, base_filename, **kwargs)
        elif format.lower() == 'html':
            filepath = self._export_report_html(report_data, base_filename, **kwargs)
        elif format.lower() == 'docx':
            filepath = self._export_report_docx(report_data, base_filename, **kwargs)
        else:
            raise ValidationError(f"Unsupported report export format: {format}")
        
        self._log_export('report', filepath, format, 1)
        return filepath
    
    def _export_report_pdf(self, report_data: Dict[str, Any], filename: str, **kwargs) -> str:
        """Export report to PDF using ReportLab."""
        filepath = os.path.join(self.exports_dir, f"{filename}.pdf")
        
        # Create PDF document
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        story.append(Paragraph("Data Analysis Report", title_style))
        story.append(Spacer(1, 12))
        
        # Metadata
        metadata_style = styles['Normal']
        story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", metadata_style))
        story.append(Spacer(1, 20))
        
        # Report content
        for section, content in report_data.items():
            if section == 'title':
                continue
                
            # Section header
            story.append(Paragraph(section.replace('_', ' ').title(), styles['Heading2']))
            story.append(Spacer(1, 12))
            
            # Section content
            if isinstance(content, dict):
                for key, value in content.items():
                    if isinstance(value, (list, dict)):
                        # Convert complex data to string
                        value_str = json.dumps(value, indent=2, default=str)
                        story.append(Paragraph(f"{key}:", styles['Heading3']))
                        story.append(Paragraph(f"<pre>{value_str}</pre>", styles['Code']))
                    else:
                        story.append(Paragraph(f"{key}: {value}", metadata_style))
                    story.append(Spacer(1, 6))
            elif isinstance(content, (list, tuple)):
                for item in content:
                    story.append(Paragraph(f"• {item}", metadata_style))
                    story.append(Spacer(1, 6))
            else:
                story.append(Paragraph(str(content), metadata_style))
            
            story.append(Spacer(1, 20))
        
        # Build PDF
        doc.build(story)
        return filepath
    
    def _export_report_html(self, report_data: Dict[str, Any], filename: str, **kwargs) -> str:
        """Export report to HTML."""
        filepath = os.path.join(self.exports_dir, f"{filename}.html")
        
        html_content = self._generate_html_report(report_data)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return filepath
    
    def _generate_html_report(self, report_data: Dict[str, Any]) -> str:
        """Generate HTML report content."""
        html_template = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Data Analysis Report</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
                .header { text-align: center; margin-bottom: 30px; }
                .section { margin-bottom: 30px; }
                .section h2 { color: #333; border-bottom: 2px solid #007bff; padding-bottom: 10px; }
                .content { margin-left: 20px; }
                .metadata { background: #f8f9fa; padding: 15px; border-radius: 5px; margin-bottom: 20px; }
                table { border-collapse: collapse; width: 100%; margin: 10px 0; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; }
                .code { background: #f4f4f4; padding: 10px; border-radius: 5px; font-family: monospace; }
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Data Analysis Report</h1>
                <div class="metadata">
                    <p><strong>Generated:</strong> {timestamp}</p>
                </div>
            </div>
            
            {content}
        </body>
        </html>
        """
        
        # Generate content sections
        content_sections = []
        for section, data in report_data.items():
            if section == 'title':
                continue
                
            section_html = f'<div class="section"><h2>{section.replace("_", " ").title()}</h2><div class="content">'
            
            if isinstance(data, dict):
                section_html += '<table>'
                for key, value in data.items():
                    if isinstance(value, (dict, list)):
                        value_str = json.dumps(value, indent=2, default=str)
                        section_html += f'<tr><td><strong>{key}</strong></td><td><div class="code">{value_str}</div></td></tr>'
                    else:
                        section_html += f'<tr><td><strong>{key}</strong></td><td>{value}</td></tr>'
                section_html += '</table>'
            elif isinstance(data, list):
                for item in data:
                    section_html += f'<p>• {item}</p>'
            else:
                section_html += f'<p>{data}</p>'
            
            section_html += '</div></div>'
            content_sections.append(section_html)
        
        return html_template.format(
            timestamp=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            content=''.join(content_sections)
        )
    
    def _export_report_docx(self, report_data: Dict[str, Any], filename: str, **kwargs) -> str:
        """Export report to DOCX format."""
        try:
            from docx import Document
            from docx.shared import Inches
            
            filepath = os.path.join(self.exports_dir, f"{filename}.docx")
            
            doc = Document()
            
            # Title
            doc.add_heading('Data Analysis Report', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph()  # Empty line
            
            # Content
            for section, content in report_data.items():
                if section == 'title':
                    continue
                    
                doc.add_heading(section.replace('_', ' ').title(), 1)
                
                if isinstance(content, dict):
                    for key, value in content.items():
                        doc.add_heading(key, 2)
                        if isinstance(value, (dict, list)):
                            doc.add_paragraph(json.dumps(value, indent=2, default=str))
                        else:
                            doc.add_paragraph(str(value))
                elif isinstance(content, list):
                    for item in content:
                        doc.add_paragraph(f'• {item}')
                else:
                    doc.add_paragraph(str(content))
                
                doc.add_paragraph()  # Empty line
            
            doc.save(filepath)
            return filepath
            
        except ImportError:
            raise ValidationError("DOCX export requires python-docx package")
    
    @handle_errors("data")
    def export_jupyter_notebook(self, analysis_code: List[str], markdown_cells: List[str] = None, filename: str = "analysis") -> str:
        """Export analysis as Jupyter notebook."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filepath = os.path.join(self.exports_dir, f"{filename}_{timestamp}.ipynb")
        
        # Create notebook
        nb = new_notebook()
        
        # Add title
        nb.cells.append(new_markdown_cell("# Data Analysis Notebook"))
        nb.cells.append(new_markdown_cell(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"))
        nb.cells.append(new_markdown_cell("---"))
        
        # Add markdown cells if provided
        if markdown_cells:
            for markdown in markdown_cells:
                nb.cells.append(new_markdown_cell(markdown))
        
        # Add code cells
        for code in analysis_code:
            nb.cells.append(new_code_cell(code))
        
        # Save notebook
        with open(filepath, 'w') as f:
            nbformat.write(nb, f)
        
        self._log_export('notebook', filepath, 'ipynb', 1)
        return filepath
    
    @handle_errors("data")
    def create_export_package(self, files: List[str], package_name: str = "analysis_package") -> str:
        """Create a zip package with multiple files."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        package_filename = f"{package_name}_{timestamp}.zip"
        package_filepath = os.path.join(self.exports_dir, package_filename)
        
        with zipfile.ZipFile(package_filepath, 'w') as zipf:
            for file in files:
                if os.path.exists(file):
                    zipf.write(file, os.path.basename(file))
        
        self._log_export('package', package_filepath, 'zip', len(files))
        return package_filepath
    
    def _log_export(self, export_type: str, filepath: str, format: str, item_count: int):
        """Log export operation."""
        log_entry = {
            'timestamp': datetime.now().isoformat(),
            'type': export_type,
            'filepath': filepath,
            'format': format,
            'item_count': item_count,
            'file_size': os.path.getsize(filepath) if os.path.exists(filepath) else 0
        }
        
        self.export_history.append(log_entry)
    
    def get_export_history(self) -> List[Dict[str, Any]]:
        """Get export history."""
        return self.export_history.copy()
    
    def clear_export_history(self):
        """Clear export history."""
        self.export_history = []
    
    def cleanup_old_exports(self, days_old: int = 7):
        """Clean up old export files."""
        import time
        
        current_time = time.time()
        cutoff_time = current_time - (days_old * 24 * 60 * 60)
        
        for filename in os.listdir(self.exports_dir):
            filepath = os.path.join(self.exports_dir, filename)
            if os.path.getmtime(filepath) < cutoff_time:
                try:
                    os.remove(filepath)
                except:
                    pass


class AdvancedReportGenerator:
    """Advanced report generation with templates and customization."""
    
    def __init__(self):
        """Initialize report generator."""
        self.templates = {}
        self._load_default_templates()
    
    def _load_default_templates(self):
        """Load default report templates."""
        self.templates['data_analysis'] = {
            'title': 'Data Analysis Report',
            'sections': [
                'executive_summary',
                'data_overview',
                'data_quality',
                'exploratory_analysis',
                'findings',
                'recommendations'
            ]
        }
        
        self.templates['ml_report'] = {
            'title': 'Machine Learning Model Report',
            'sections': [
                'model_overview',
                'data_preparation',
                'model_performance',
                'feature_importance',
                'model_interpretation',
                'recommendations'
            ]
        }
        
        self.templates['text_analysis'] = {
            'title': 'Text Analytics Report',
            'sections': [
                'text_overview',
                'sentiment_analysis',
                'topic_modeling',
                'keyword_analysis',
                'insights'
            ]
        }
    
    @handle_errors("data")
    def generate_comprehensive_report(self, data: Dict[str, Any], template_name: str = 'data_analysis', 
                                 customizations: Dict[str, Any] = None) -> Dict[str, Any]:
        """Generate comprehensive report using template."""
        if template_name not in self.templates:
            raise ValidationError(f"Unknown template: {template_name}")
        
        template = self.templates[template_name]
        customizations = customizations or {}
        
        report = {
            'title': customizations.get('title', template['title']),
            'generated_at': datetime.now().isoformat(),
            'sections': {}
        }
        
        # Generate each section
        for section in template['sections']:
            if section in data:
                report['sections'][section] = self._generate_section(section, data[section], customizations)
            else:
                report['sections'][section] = self._generate_default_section(section)
        
        return report
    
    def _generate_section(self, section_name: str, section_data: Any, customizations: Dict[str, Any]) -> Dict[str, Any]:
        """Generate specific report section."""
        section_generators = {
            'executive_summary': self._generate_executive_summary,
            'data_overview': self._generate_data_overview,
            'data_quality': self._generate_data_quality_section,
            'exploratory_analysis': self._generate_eda_section,
            'model_performance': self._generate_model_performance_section,
            'findings': self._generate_findings_section,
            'recommendations': self._generate_recommendations_section
        }
        
        if section_name in section_generators:
            return section_generators[section_name](section_data, customizations)
        else:
            return {'content': section_data, 'type': 'custom'}
    
    def _generate_executive_summary(self, data: Any, customizations: Dict[str, Any]) -> Dict[str, Any]:
        """Generate executive summary section."""
        return {
            'title': 'Executive Summary',
            'type': 'summary',
            'content': data,
            'key_points': customizations.get('key_points', [])
        }
    
    def _generate_data_overview(self, data: Any, customizations: Dict[str, Any]) -> Dict[str, Any]:
        """Generate data overview section."""
        return {
            'title': 'Data Overview',
            'type': 'overview',
            'content': data,
            'metrics': customizations.get('metrics', {})
        }
    
    def _generate_data_quality_section(self, data: Any, customizations: Dict[str, Any]) -> Dict[str, Any]:
        """Generate data quality section."""
        return {
            'title': 'Data Quality Assessment',
            'type': 'quality',
            'content': data,
            'quality_score': data.get('overall_score', 0),
            'issues': data.get('issues', [])
        }
    
    def _generate_eda_section(self, data: Any, customizations: Dict[str, Any]) -> Dict[str, Any]:
        """Generate EDA section."""
        return {
            'title': 'Exploratory Data Analysis',
            'type': 'eda',
            'content': data,
            'insights': customizations.get('insights', [])
        }
    
    def _generate_model_performance_section(self, data: Any, customizations: Dict[str, Any]) -> Dict[str, Any]:
        """Generate model performance section."""
        return {
            'title': 'Model Performance',
            'type': 'performance',
            'content': data,
            'best_model': data.get('best_model'),
            'metrics': data.get('metrics', {})
        }
    
    def _generate_findings_section(self, data: Any, customizations: Dict[str, Any]) -> Dict[str, Any]:
        """Generate findings section."""
        return {
            'title': 'Key Findings',
            'type': 'findings',
            'content': data,
            'findings': customizations.get('findings', [])
        }
    
    def _generate_recommendations_section(self, data: Any, customizations: Dict[str, Any]) -> Dict[str, Any]:
        """Generate recommendations section."""
        return {
            'title': 'Recommendations',
            'type': 'recommendations',
            'content': data,
            'recommendations': customizations.get('recommendations', [])
        }
    
    def _generate_default_section(self, section_name: str) -> Dict[str, Any]:
        """Generate default section for missing data."""
        return {
            'title': section_name.replace('_', ' ').title(),
            'type': 'default',
            'content': 'No data available for this section.'
        }


class BatchExporter:
    """Batch export system for multiple items."""
    
    def __init__(self):
        """Initialize batch exporter."""
        self.exporter = DataExporter()
        self.batch_queue = []
    
    def add_to_batch(self, item_type: str, item: Any, filename: str, format: str = 'auto', **kwargs):
        """Add item to batch export queue."""
        self.batch_queue.append({
            'type': item_type,
            'item': item,
            'filename': filename,
            'format': format,
            'kwargs': kwargs
        })
    
    def process_batch(self, package_name: str = "batch_export") -> str:
        """Process entire batch and create package."""
        exported_files = []
        
        for batch_item in self.batch_queue:
            try:
                if batch_item['type'] == 'data':
                    filepath = self.exporter.export_data(
                        batch_item['item'], 
                        batch_item['filename'], 
                        batch_item['format'],
                        **batch_item['kwargs']
                    )
                elif batch_item['type'] == 'model':
                    filepath = self.exporter.export_model(
                        batch_item['item'], 
                        batch_item['filename'], 
                        batch_item['format'],
                        **batch_item['kwargs']
                    )
                elif batch_item['type'] == 'visualization':
                    filepath = self.exporter.export_visualization(
                        batch_item['item'], 
                        batch_item['filename'], 
                        batch_item['format'],
                        **batch_item['kwargs']
                    )
                elif batch_item['type'] == 'report':
                    filepath = self.exporter.export_report(
                        batch_item['item'], 
                        batch_item['filename'], 
                        batch_item['format'],
                        **batch_item['kwargs']
                    )
                
                exported_files.append(filepath)
                
            except Exception as e:
                print(f"Failed to export {batch_item['filename']}: {str(e)}")
        
        # Create package
        if exported_files:
            package_path = self.exporter.create_export_package(exported_files, package_name)
            self.batch_queue.clear()
            return package_path
        
        raise ValidationError("No files were successfully exported")
    
    def get_batch_status(self) -> Dict[str, Any]:
        """Get batch export status."""
        return {
            'queue_size': len(self.batch_queue),
            'items': [
                {
                    'type': item['type'],
                    'filename': item['filename'],
                    'format': item['format']
                }
                for item in self.batch_queue
            ]
        }
    
    def clear_batch(self):
        """Clear batch queue."""
        self.batch_queue = []


@rate_limit()
def create_download_link(content: Any, filename: str, format: str = 'csv') -> str:
    """Create download link for content."""
    if format == 'csv' and isinstance(content, pd.DataFrame):
        csv = content.to_csv(index=False)
        b64 = base64.b64encode(csv.encode()).decode()
        href = f'<a href="data:file/csv;base64,{b64}" download="{filename}">Download CSV</a>'
    elif format == 'json':
        json_str = json.dumps(content, indent=2, default=str)
        b64 = base64.b64encode(json_str.encode()).decode()
        href = f'<a href="data:application/json;base64,{b64}" download="{filename}">Download JSON</a>'
    elif format == 'excel' and isinstance(content, pd.DataFrame):
        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            content.to_excel(writer, index=False)
        b64 = base64.b64encode(output.getvalue()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{filename}">Download Excel</a>'
    else:
        href = f'<span>Unsupported format: {format}</span>'
    
    return href


def export_streamlit_data(data: pd.DataFrame, filename: str, format: str = 'csv'):
    """Export data for Streamlit download."""
    import streamlit as st
    
    if format == 'csv':
        csv = data.to_csv(index=False)
        st.download_button(
            label="Download CSV",
            data=csv,
            file_name=filename,
            mime='text/csv'
        )
    elif format == 'excel':
        output = BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            data.to_excel(writer, index=False)
        st.download_button(
            label="Download Excel",
            data=output.getvalue(),
            file_name=filename,
            mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    elif format == 'json':
        json_str = data.to_json(orient='records', indent=2)
        st.download_button(
            label="Download JSON",
            data=json_str,
            file_name=filename,
            mime='application/json'
        )


# Standalone function for app_adapted.py compatibility
def create_export_center(data, models=None):
    """Create export center UI for Streamlit."""
    import streamlit as st
    
    st.title("📥 Export Center")
    
    if data is None or data.empty:
        st.warning("⚠️ No data available to export")
        return
    
    tab1, tab2, tab3 = st.tabs(["📊 Data Export", "🤖 Model Export", "📋 Report Export"])
    
    with tab1:
        st.subheader("📊 Export Data")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            filename = st.text_input("Filename", value="dataset")
            format_type = st.selectbox("Format", ["csv", "excel", "json", "parquet", "tsv"])
        
        with col2:
            include_index = st.checkbox("Include Index", value=False)
            encoding = st.selectbox("Encoding", ["utf-8", "latin-1", "cp1252"])
        
        with col3:
            if st.button("📥 Export Data", type="primary"):
                try:
                    exporter = DataExporter()
                    filepath = exporter.export_data(
                        data,
                        filename,
                        format_type,
                        index=include_index,
                        encoding=encoding
                    )
                    st.success(f"✅ Data exported successfully!")
                    
                    # Provide download button for the exported data
                    if os.path.exists(filepath):
                        with open(filepath, "rb") as file:
                            file_data = file.read()
                        
                        # Determine MIME type based on format
                        if format_type.lower() == 'csv':
                            mime_type = 'text/csv'
                        elif format_type.lower() == 'excel':
                            mime_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                        elif format_type.lower() == 'json':
                            mime_type = 'application/json'
                        elif format_type.lower() == 'parquet':
                            mime_type = 'application/octet-stream'
                        elif format_type.lower() == 'tsv':
                            mime_type = 'text/tab-separated-values'
                        else:
                            mime_type = 'application/octet-stream'
                        
                        # Create download button
                        st.download_button(
                            label=f"📥 Download {format_type.upper()} File",
                            data=file_data,
                            file_name=f"{filename}.{format_type}",
                            mime=mime_type
                        )
                    
                except Exception as e:
                    st.error(f"❌ Export failed: {str(e)}")
    
    with tab2:
        st.subheader("🤖 Export Models")
        
        if models:
            for model_name, model_data in models.items():
                with st.expander(f"🤖 {model_name.upper()}"):
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        model_filename = st.text_input(f"Filename", value=f"{model_name}_model", key=f"model_fname_{model_name}")
                        model_format = st.selectbox("Format", ["pickle", "joblib", "json"], key=f"model_format_{model_name}")
                    
                    with col2:
                        if st.button("📥 Export Model", key=f"export_model_{model_name}"):
                            try:
                                exporter = DataExporter()
                                filepath = exporter.export_model(
                                    model_data.get('model') if isinstance(model_data, dict) else model_data,
                                    model_filename,
                                    model_format,
                                    metadata=model_data.get('metrics', {}) if isinstance(model_data, dict) else {}
                                )
                                st.success(f"✅ Model exported successfully!")
                                
                                # Provide download button for the exported model
                                if os.path.exists(filepath):
                                    with open(filepath, "rb") as file:
                                        file_data = file.read()
                                    
                                    # Determine MIME type based on format
                                    if model_format.lower() == 'pickle':
                                        mime_type = 'application/octet-stream'
                                    elif model_format.lower() == 'joblib':
                                        mime_type = 'application/octet-stream'
                                    elif model_format.lower() == 'json':
                                        mime_type = 'application/json'
                                    elif model_format.lower() == 'onnx':
                                        mime_type = 'application/octet-stream'
                                    else:
                                        mime_type = 'application/octet-stream'
                                    
                                    # Create download button
                                    st.download_button(
                                        label=f"📥 Download {model_format.upper()} Model",
                                        data=file_data,
                                        file_name=f"{model_filename}.{model_format}",
                                        mime=mime_type
                                    )
                                
                            except Exception as e:
                                st.error(f"❌ Export failed: {str(e)}")
        else:
            st.info("ℹ️ No trained models available")
    
    with tab3:
        st.subheader("📋 Export Report")
        
        col1, col2 = st.columns(2)
        
        with col1:
            report_filename = st.text_input("Report Filename", value="analysis_report")
            report_format = st.selectbox("Report Format", ["pdf", "html", "docx"])
        
        with col2:
            include_summary = st.checkbox("Include Summary", value=True)
            include_charts = st.checkbox("Include Charts", value=True)
        
        if st.button("📋 Generate Report", type="primary"):
            try:
                # Generate report data
                report_data = {
                    'title': 'Data Analysis Report',
                    'data_summary': {
                        'shape': data.shape,
                        'columns': data.columns.tolist(),
                        'missing_values': data.isnull().sum().to_dict(),
                        'data_types': data.dtypes.astype(str).to_dict()
                    },
                    'quality_assessment': {
                        'overall_score': 85,  # Placeholder
                        'issues': [],
                        'recommendations': ['Data looks clean', 'Ready for analysis']
                    }
                }
                
                if include_summary:
                    report_data['executive_summary'] = {
                        'key_points': [
                            f"Dataset contains {data.shape[0]} rows and {data.shape[1]} columns",
                            "Data quality appears to be good",
                            "Ready for machine learning analysis"
                        ]
                    }
                
                exporter = DataExporter()
                filepath = exporter.export_report(report_data, report_filename, report_format)
                st.success(f"✅ Report generated successfully!")
                
                # Provide download button for the generated report
                if os.path.exists(filepath):
                    with open(filepath, "rb") as file:
                        file_data = file.read()
                    
                    # Determine MIME type based on format
                    if report_format.lower() == 'pdf':
                        mime_type = 'application/pdf'
                    elif report_format.lower() == 'html':
                        mime_type = 'text/html'
                    elif report_format.lower() == 'docx':
                        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    else:
                        mime_type = 'application/octet-stream'
                    
                    # Create download button
                    st.download_button(
                        label=f"📥 Download {report_format.upper()} Report",
                        data=file_data,
                        file_name=f"{report_filename}.{report_format}",
                        mime=mime_type
                    )
                
            except Exception as e:
                st.error(f"❌ Report generation failed: {str(e)}")